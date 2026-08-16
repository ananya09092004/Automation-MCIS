"""Read/write documents across common formats for AI-assisted editing.

Plain text (str) is the universal exchange format: read_document always
returns a str, write_document always takes a str. Each format's library
only handles the conversion in/out at the edges.

PDF is read-only here -- rewriting an arbitrary PDF's layout reliably
isn't something simple tools can do safely, so an edited PDF is saved
as a new .docx next to the original instead of corrupting/guessing at
the original's layout.
"""
from pathlib import Path


def read_document(path: str) -> str:
    ext = Path(path).suffix.lower()

    if ext in (".txt", ".md"):
        return Path(path).read_text(encoding="utf-8", errors="replace")

    if ext == ".docx":
        import docx
        doc = docx.Document(path)
        return "\n".join(paragraph.text for paragraph in doc.paragraphs)

    if ext == ".pdf":
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            return "\n".join(page.extract_text() or "" for page in pdf.pages)

    if ext in (".xlsx", ".xls"):
        import openpyxl
        workbook = openpyxl.load_workbook(path, data_only=True)
        lines = []
        for sheet in workbook.worksheets:
            lines.append(f"# Sheet: {sheet.title}")
            for row in sheet.iter_rows(values_only=True):
                lines.append(",".join("" if cell is None else str(cell) for cell in row))
        return "\n".join(lines)

    raise ValueError(f"Unsupported file type for reading: {ext or '(no extension)'}")


def write_document(path: str, content: str) -> None:
    ext = Path(path).suffix.lower()

    if ext in (".txt", ".md"):
        Path(path).write_text(content, encoding="utf-8")
        return

    if ext == ".docx":
        import docx
        doc = docx.Document()
        for line in content.split("\n"):
            doc.add_paragraph(line)
        doc.save(path)
        return

    if ext == ".pdf":
        new_path = str(Path(path).with_suffix("")) + ".edited.docx"
        write_document(new_path, content)
        raise ValueError(
            f"Can't safely rewrite a PDF in place -- saved the edited "
            f"version as {new_path} instead."
        )

    if ext in (".xlsx", ".xls"):
        import openpyxl
        workbook = openpyxl.Workbook()
        sheet = workbook.active
        for row_index, line in enumerate(content.split("\n"), start=1):
            for col_index, cell_value in enumerate(line.split(","), start=1):
                sheet.cell(row=row_index, column=col_index, value=cell_value)
        workbook.save(path)
        return

    raise ValueError(f"Unsupported file type for writing: {ext or '(no extension)'}")
